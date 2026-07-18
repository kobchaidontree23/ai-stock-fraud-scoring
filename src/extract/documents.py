"""Extract qualitative text + heuristic red-flag signals from filings.

Reads the auditor's report and notes (DOCX / legacy DOC / PDF) and derives
cheap Thai-keyword flags (auditor opinion type, going-concern, key audit
matters, related-party prominence).  The raw text is also returned so the LLM
layer can reason over it; the heuristics give signal even with no API key.
"""
from __future__ import annotations
import os
import glob
import re


# ---------------------------------------------------------------------------
# text extraction by format
# ---------------------------------------------------------------------------
def _from_docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def _from_pdf(path: str, max_pages: int = 12) -> str:
    import pdfplumber
    out = []
    with pdfplumber.open(path) as pdf:
        for pg in pdf.pages[:max_pages]:
            out.append(pg.extract_text() or "")
    return "\n".join(out)


def _from_doc(path: str) -> str:
    """Legacy binary .doc via Word COM (Windows); empty string if unavailable."""
    try:
        import win32com.client as win32
        import pythoncom
        pythoncom.CoInitialize()
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
            text = doc.Content.Text
            doc.Close(False)
            return text
        finally:
            word.Quit()
    except Exception:
        return ""


def get_text(path: str, max_pages: int = 12) -> str:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            return _from_docx(path)
        if ext == ".doc":
            return _from_doc(path)
        if ext == ".pdf":
            return _from_pdf(path, max_pages=max_pages)
    except Exception:
        return ""
    return ""


# ---------------------------------------------------------------------------
# Thai heuristic signals
# ---------------------------------------------------------------------------
# NOTE: "อย่างไม่มีเงื่อนไข" (unqualified) contains "มีเงื่อนไข", so the
# qualified pattern must require the "อย่างมีเงื่อนไข"/"แบบมีเงื่อนไข" form
# (no bare "มีเงื่อนไข") to avoid matching the negated clean opinion.
OPINION_PATTERNS = [
    ("disclaimer", r"ไม่แสดงความเห็น", 100),
    ("adverse", r"ความเห็นในทางตรงข้าม|ความเห็นที่เป็นปฏิเสธ", 95),
    # Only the *Basis for Qualified Opinion* section heading is a reliable
    # signal. Bare "แสดงความเห็นอย่างมีเงื่อนไข" also appears in KAM boilerplate
    # ("ข้าพเจ้าไม่ได้แสดงความเห็นอย่างมีเงื่อนไขต่อกรณีเหล่านี้" = "I do NOT
    # express a qualified opinion on these matters") — a clean report's basis
    # heading is "เกณฑ์ในการแสดงความเห็น" without "อย่างมีเงื่อนไข".
    ("qualified", r"เกณฑ์ในการแสดงความเห็นอย่างมีเงื่อนไข|เกณฑ์ในการแสดงความเห็นแบบมีเงื่อนไข", 70),
    ("unqualified", r"อย่างไม่มีเงื่อนไข|ถูกต้องตามที่ควรในสาระสำคัญ|แสดงความเห็นว่างบการเงิน", 10),
]
GOING_CONCERN = r"ความไม่แน่นอน[^\n]{0,40}ดำเนินงานต่อเนื่อง|อาจมีข้อสงสัย[^\n]{0,20}ดำเนินงานต่อเนื่อง|ความสามารถในการดำเนินงานต่อเนื่อง"
EMPHASIS = r"เน้นข้อมูลและเหตุการณ์|การเน้นข้อมูล|ข้อมูลและเหตุการณ์ที่เน้น"
KAM = r"เรื่องสำคัญในการตรวจสอบ|เรื่องที่สำคัญในการตรวจ"
RELATED = r"ที่เกี่ยวข้องกัน|บุคคลหรือกิจการที่เกี่ยวข้อง|รายการธุรกิจกับกิจการที่เกี่ยวข้อง"

# audit firms -> (display name, is_big4). A switch from Big-4 to a small local
# firm is a recognised pre-distress governance red flag.
AUDIT_FIRMS = [
    ("EY", ["อีวาย"], True),
    ("PwC", ["ไพร้ซวอเตอร์", "ไพร้ซ"], True),
    ("KPMG", ["เคพีเอ็มจี", "เคพีเอ็ม"], True),
    ("Deloitte", ["ดีลอยท์", "ดีลอยต์"], True),
    ("Dharmniti", ["ธรรมนิติ"], False),
    ("GrantThornton", ["แกรนท์", "กรินทร์"], False),
    ("Mazars", ["มาซาร์"], False),
    ("BDO", ["บีดีโอ"], False),
]


def _audit_firm(text: str):
    for name, keys, big4 in AUDIT_FIRMS:
        if any(k in text for k in keys):
            return name, big4
    return None, None


def _audit_license(text: str):
    m = (re.search(r"ผู้สอบบัญชีรับอนุญาต\D{0,25}?(\d{3,5})", text)
         or re.search(r"เลขทะเบียน\D{0,12}(\d{3,5})", text))
    return m.group(1) if m else None


def analyze_auditor(text: str) -> dict:
    if not text:
        return {"available": False}
    opinion, op_score = "unknown", 40
    # most severe present wins
    for name, pat, sc in OPINION_PATTERNS:
        if re.search(pat, text):
            opinion, op_score = name, sc
            if name in ("disclaimer", "adverse", "qualified"):
                break
    gc = re.search(GOING_CONCERN, text)
    emph = re.search(EMPHASIS, text)
    kam = len(re.findall(KAM, text))

    def _snip(m, span=280):
        s = re.sub(r"\s+", " ", text[m.start():m.start() + span]).strip()
        return s

    score = op_score
    if gc:
        score = max(score, 80)
    if emph and opinion == "unqualified":
        score = max(score, 50)
    firm, big4 = _audit_firm(text)
    return {"available": True, "opinion": opinion, "going_concern": bool(gc),
            "emphasis_of_matter": bool(emph), "has_kam": kam > 0,
            "emphasis_text": _snip(emph) if emph else None,
            "going_concern_text": _snip(gc) if gc else None,
            "auditor_license": _audit_license(text),
            "auditor_firm": firm, "auditor_big4": big4,
            "score": score, "band": None,
            "detail": f"opinion={opinion}"
            + (", GOING-CONCERN uncertainty" if gc else "")
            + (", emphasis-of-matter" if emph else "")}


def analyze_related_party(text: str) -> dict:
    if not text:
        return {"available": False}
    n = len(re.findall(RELATED, text))
    return {"available": True, "related_party_mentions": n,
            "detail": f"{n} related-party references in notes"}


# ---------------------------------------------------------------------------
# per-year assembly
# ---------------------------------------------------------------------------
def _year_dir(base: str, year: int):
    """Find the year folder holding the statements/auditor report for a ticker.

    Works across "financial statement" / "financial report" naming and year
    folders like "JKN2565" or "DELTA 2565"; prefers a folder that actually
    contains an auditor report or financial statements over a one-report match.
    """
    for yy in (year, year - 543):          # folder may use BE (25xx) or CE (20xx)
        cands = glob.glob(os.path.join(base, "*", f"*{yy}*"))
        good = [c for c in cands if os.path.isdir(c) and (
            glob.glob(os.path.join(c, "AUDITOR_REPORT.*"))
            or glob.glob(os.path.join(c, "FINANCIAL_STATEMENTS.*")))]
        if good:
            return good[0]
    return None


# One Report (56-1) — the narrative filing: business, MD&A, corporate
# governance, board/management, subsidiaries. Governance-relevant sections.
OR_PREFER = ["MANAGEMENT_AND_CG", "INFO_DIRECTOR_MANAGEMENT",
             "INFO_SUBSIDIARIES_EXECUTIVES", "INFO_CHIEF_INTERNAL_AUDIT",
             "BUSINESS"]


def _one_report_dir(ticker_dir: str, year: int):
    for yy in (year, year - 543):          # BE (25xx) or CE (20xx) folder naming
        cands = glob.glob(os.path.join(ticker_dir, "one report", f"*{yy}*"))
        if cands:
            return cands[0]
    return None


def extract_one_report(ticker_dir: str, year: int, use_doc_com: bool = True,
                       max_chars: int = 14000, combined_pages: int = 25) -> dict:
    """Pull governance-relevant One Report text + related-party density.

    Prefers the dedicated section files (older filings split the 56-1 into
    MANAGEMENT_AND_CG, INFO_DIRECTOR_MANAGEMENT, ...). Newer filings ship one
    combined PDF; those are read with a bounded page window (coverage is
    partial by design — the auditor-rotation governance signal does not depend
    on this text). Degrades cleanly on any parse failure.
    """
    d = _one_report_dir(ticker_dir, year)
    if not d:
        return {"available": False}
    picks = []
    for pat in OR_PREFER:
        picks += sorted(glob.glob(os.path.join(d, f"{pat}*")))
    combined = False
    if not picks:
        combined = True
        allf = [f for f in glob.glob(os.path.join(d, "*"))
                if f.lower().endswith((".pdf", ".doc", ".docx"))
                and "STRUCTURE" not in os.path.basename(f).upper()]
        allf.sort(key=lambda p: os.path.getsize(p), reverse=True)
        picks = allf[:1]

    text, used = "", []
    for f in picks:
        if len(text) >= max_chars:
            break
        if f.lower().endswith(".doc") and not use_doc_com:
            continue
        t = get_text(f, max_pages=(combined_pages if combined else 12))
        if t:
            text += "\n" + t
            used.append(os.path.basename(f))
    text = text[:max_chars]
    return {"available": bool(text), "text": text,
            "related_party_mentions": len(re.findall(RELATED, text)),
            "combined": combined, "sources": used}


def extract_documents(ticker_dir: str, years, use_doc_com: bool = True,
                      max_chars: int = 16000, with_one_report: bool = True) -> dict:
    """Return per-year auditor / notes / one-report signals."""
    out = {}
    for y in years:
        fdir = _year_dir(ticker_dir, y)
        rec = {"auditor_text": "", "auditor": {"available": False},
               "related_party": {"available": False}, "notes_excerpt": "",
               "one_report": {"available": False}}
        if fdir:
            aud = (glob.glob(os.path.join(fdir, "AUDITOR_REPORT.*")) or [None])[0]
            notes = (glob.glob(os.path.join(fdir, "NOTES.*")) or [None])[0]
            if aud and (use_doc_com or not aud.lower().endswith(".doc")):
                t = get_text(aud)                     # analyse FULL text: the
                rec["auditor"] = analyze_auditor(t)   # firm/licence sits in the
                rec["auditor_text"] = t[:max_chars]   # end-of-report signature

            if notes and (use_doc_com or not notes.lower().endswith(".doc")):
                nt = get_text(notes)
                rec["related_party"] = analyze_related_party(nt)
                rec["notes_excerpt"] = nt[:max_chars]
        if with_one_report:
            rec["one_report"] = extract_one_report(ticker_dir, y, use_doc_com)
        out[y] = rec
    return out


if __name__ == "__main__":
    import sys, json
    d = extract_documents(sys.argv[1], [int(sys.argv[2])])
    for y, r in d.items():
        r2 = {k: (v if k not in ("auditor_text", "notes_excerpt") else f"<{len(v)} chars>")
              for k, v in r.items()}
        print(y, json.dumps(r2, ensure_ascii=False, indent=2))
